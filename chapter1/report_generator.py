from docx import Document
from docx.shared import RGBColor
from chatgpt_client import ChatGPTClient
from myutils import read_yaml, read_text_file


class ReportGenerator:
    """
    This is the main class for generating the reports related to the project.
    """

    def __init__(self, template_config_path: str):
        self.chatgpt_client = ChatGPTClient()
        self.project_report = Document()
        self.prompt_eng_report = Document()
        self.template_config_path = template_config_path
        self.template = self.load_template_config()
        self.sections = self.template["content"]
        self.textfiles_dir = self.template["textfiles_dir"]


    def generate_content_menu(self, sections):
        for section in sections:
            indent = (section['level'] - 1) * 4
            self.project_report.add_heading((indent * ' ' + section['title']), level=section['level'])
            if len(section['subsections']):
                self.generate_content_menu(section['subsections'])
            else:
                continue

    def write_the_chapter(self, sections):
        for section in sections:
            # write the title of the section
            print("writing section title: " + section['title'] + " ...")
            self.project_report.add_heading(section['title'], level=section['level'])
            self.prompt_eng_report.add_heading(section['title'], level=section['level'])
            # write the text file to the section
            if len(section['text_files']):
                for text_file in section['text_files']:
                    print("writing text file: " + text_file + " ...")
                    text = read_text_file(self.textfiles_dir + text_file)
                    self.project_report.add_paragraph(text)
                    self.prompt_eng_report.add_paragraph(text)
            # write the prompt to the prompt_eng section
            if len(section['prompt_files']):
                for prompt_file in section['prompt_files']:
                    print("reading prompt file: " + prompt_file + " ...")
                    prompt = read_text_file(self.textfiles_dir + prompt_file)
                    p1 = self.prompt_eng_report.add_paragraph()
                    p1.add_run(prompt).font.color.rgb = RGBColor(26, 93, 26)  # green
                    # send the prompt to the chatgpt
                    print("sending prompt to chatgpt ...")
                    response_to_prompt_from_file = self.chatgpt_client.get_completion_from_messages(prompt)
                    # write the response to the section
                    print("writing response to prompt to prompt eng report...")
                    p2 = self.prompt_eng_report.add_paragraph()
                    p2.add_run(response_to_prompt_from_file).font.color.rgb = RGBColor(128, 0, 128)  # purple
                    # send the prompt to chatgpt to analyze
                    print("sending prompt to chatgpt to analyze ...")
                    response_to_analyze_prompt = self.chatgpt_client.analyze_prompt(prompt)
                    # write the response to the section
                    print("writing response to analyze prompt to prompt eng report...")
                    p3 = self.prompt_eng_report.add_paragraph()
                    p3.add_run(response_to_analyze_prompt).font.color.rgb = RGBColor(0, 0, 0)  # black
                    # send the prompt to chatgpt to get modified prompt
                    print("sending prompt to chatgpt to modify ...")
                    ai_modified_prompt = self.chatgpt_client.modify_prompt(prompt)
                    # write the ai_modified_prompt to the section
                    print("writing ai modified prompt to prompt eng report...")
                    p4 = self.prompt_eng_report.add_paragraph()
                    p4.add_run(ai_modified_prompt).font.color.rgb = RGBColor(26, 93, 26)  # green
                    # send the ai_modified_prompt to chatgpt
                    print("sending ai modified prompt to chatgpt ...")
                    response_to_ai_modified_prompt = self.chatgpt_client.get_completion_from_messages(ai_modified_prompt)
                    # write the response to the section
                    print("writing response to ai modified prompt to prompt eng report...")
                    p5 = self.prompt_eng_report.add_paragraph()
                    p5.add_run(response_to_ai_modified_prompt).font.color.rgb = RGBColor(128, 0, 128)  # purple
                    # write the final answer to the project report section
                    print("writing final answer to project report section...")
                    p6 = self.project_report.add_paragraph()
                    p6.add_run(response_to_ai_modified_prompt).font.color.rgb = RGBColor(128, 0, 128)  # purple

            if len(section['subsections']):
                self.write_the_chapter(section['subsections'])

    def load_template_config(self):
        return read_yaml(self.template_config_path)


    def save_document(self):
        self.project_report.save('output/project_report.docx')
        self.prompt_eng_report.save('output/prompt_eng_report.docx')

